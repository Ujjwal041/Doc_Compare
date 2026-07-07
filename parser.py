"""
Document Parser — extracts structured text from Word (.docx) files.
Returns sections dict: {section_title: [paragraphs], ...}
Also extracts tables as list of dicts.
"""

from docx import Document
from pathlib import Path


def parse_docx(file_path: str) -> dict:
    """
    Parse a .docx file into structured sections.
    Returns:
        {
            "full_text": str,
            "sections": {"section_title": "content", ...},
            "tables": [{"headers": [...], "rows": [[...], ...]}, ...],
            "metadata": {"filename": str, "total_paragraphs": int}
        }
    """
    doc = Document(file_path)
    filename = Path(file_path).name

    full_text_parts = []
    sections = {}
    current_section = "PREAMBLE"
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        full_text_parts.append(text)

        # Detect section headings
        style_name = (para.style.name.lower() if para.style and para.style.name else "")
        is_heading = (
            "heading" in style_name or
            text.startswith("SECTION ") or
            text.startswith("PART ") or
            text.startswith("AMENDMENT TO") or
            (text.isupper() and len(text) > 5 and len(text) < 80)
        )

        if is_heading:
            # Save previous section
            if current_content:
                sections[current_section] = "\n".join(current_content)
            current_section = text
            current_content = []
        else:
            current_content.append(text)

    # Save last section
    if current_content:
        sections[current_section] = "\n".join(current_content)

    # Extract tables
    tables = []
    for table in doc.tables:
        rows = []
        headers = []
        for i, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if i == 0:
                headers = cells
            else:
                if any(cells):  # skip empty rows
                    rows.append(cells)
        if headers:
            tables.append({"headers": headers, "rows": rows})

    return {
        "full_text": "\n".join(full_text_parts),
        "sections": sections,
        "tables": tables,
        "metadata": {
            "filename": filename,
            "total_paragraphs": len(full_text_parts),
            "total_sections": len(sections),
            "total_tables": len(tables),
        }
    }


def extract_effective_date(parsed_doc: dict) -> str:
    """
    Try to extract effective date from document metadata section.
    Returns date string like '2024-01-01' or None.
    """
    import re
    date_patterns = [
        r"Effective Date[:\s]+(\w+ \d+,\s*\d{4})",
        r"Effective[:\s]+(\w+ \d+,\s*\d{4})",
        r"effective (\w+ \d+,\s*\d{4})",
    ]
    text = parsed_doc["full_text"][:2000]  # check first 2000 chars
    for pattern in date_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            from datetime import datetime
            try:
                dt = datetime.strptime(match.group(1).strip(), "%B %d, %Y")
                return dt.strftime("%Y-%m-%d")
            except ValueError:
                pass
    return None
